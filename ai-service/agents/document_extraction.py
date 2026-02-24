"""
Document Extraction Agent — Extracts assessment structure from PDF/DOCX files.
3-node LangGraph: parse_document → extract_structure → validate_output
"""

import json
import io
import logging
from typing import TypedDict, Optional
from langgraph.graph import StateGraph, END
from config import get_llm
from prompts.document_extraction import ASSESSMENT_EXTRACTION_PROMPT

logger = logging.getLogger("document-extraction")


# ─── State ────────────────────────────────────────────

class DocumentExtractionState(TypedDict):
    file_bytes: bytes
    file_type: str  # "pdf" or "docx"
    grade_hint: Optional[str]
    # After parse
    raw_text: str
    # After extract
    extracted_json: dict
    # After validate
    validated: bool
    errors: list[str]


# ─── Node 1: Parse Document ──────────────────────────

def parse_document_node(state: DocumentExtractionState) -> dict:
    """Extract raw text from PDF or DOCX file."""
    file_bytes = state["file_bytes"]
    file_type = state["file_type"]
    raw_text = ""

    try:
        if file_type == "pdf":
            import fitz  # pymupdf
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            for page in doc:
                raw_text += page.get_text() + "\n"
            doc.close()

        elif file_type == "docx":
            from docx import Document
            doc = Document(io.BytesIO(file_bytes))
            for para in doc.paragraphs:
                raw_text += para.text + "\n"
            # Also extract tables
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    raw_text += " | ".join(cells) + "\n"

        else:
            raw_text = file_bytes.decode("utf-8", errors="ignore")

    except Exception as e:
        logger.error(f"Document parse failed: {e}")
        raw_text = f"ERROR: Could not parse document: {str(e)}"

    logger.info(f"Parsed document: {len(raw_text)} chars, type={file_type}")
    return {"raw_text": raw_text}


# ─── Node 2: Extract Structure (LLM) ─────────────────

def extract_structure_node(state: DocumentExtractionState) -> dict:
    """Use LLM to extract assessment structure from raw text."""
    raw_text = state.get("raw_text", "")

    if raw_text.startswith("ERROR:"):
        return {"extracted_json": {}, "errors": [raw_text]}

    # Truncate to avoid token limits (keep first 8000 chars)
    text_for_llm = raw_text[:8000]

    prompt = ASSESSMENT_EXTRACTION_PROMPT.format(document_text=text_for_llm)

    llm = get_llm(temperature=0.1)
    response = llm.invoke(prompt)

    # Parse JSON from response
    content = response.content.strip()
    # Strip markdown code fences if present
    if content.startswith("```"):
        lines = content.split("\n")
        content = "\n".join(lines[1:])
        if content.endswith("```"):
            content = content[:-3].strip()

    try:
        extracted = json.loads(content)
    except json.JSONDecodeError as e:
        logger.error(f"Failed to parse LLM JSON: {e}")
        return {"extracted_json": {}, "errors": [f"JSON parse error: {str(e)}"]}

    # Apply grade hint if provided
    grade_hint = state.get("grade_hint")
    if grade_hint and not extracted.get("grade"):
        extracted["grade"] = grade_hint

    return {"extracted_json": extracted}


# ─── Node 3: Validate Output ─────────────────────────

def validate_output_node(state: DocumentExtractionState) -> dict:
    """Validate the extracted assessment structure."""
    data = state.get("extracted_json", {})
    errors = list(state.get("errors", []))

    if not data:
        if not errors:
            errors.append("No data extracted from document")
        return {"validated": False, "errors": errors}

    # Validate required fields
    if not data.get("grade"):
        errors.append("Could not determine grade level")

    domain_scores = data.get("domainMaxScores", {})
    valid_domains = {"reading", "readingComp", "spelling", "numeracy", "writing"}

    if not domain_scores:
        errors.append("No domain max scores extracted")
    else:
        for domain in valid_domains:
            if domain not in domain_scores:
                errors.append(f"Missing domain: {domain}")
            elif not isinstance(domain_scores[domain], (int, float)) or domain_scores[domain] <= 0:
                errors.append(f"Invalid max score for {domain}: {domain_scores.get(domain)}")

    sections = data.get("sections", [])
    if not sections:
        errors.append("No assessment sections extracted")

    # Ensure at least reading, numeracy, and writing are present
    found_domains = set()
    for section in sections:
        found_domains.add(section.get("domain", ""))
    missing_critical = {"reading", "numeracy"} - found_domains
    if missing_critical:
        errors.append(f"Missing critical domains: {', '.join(missing_critical)}")

    return {"validated": len(errors) == 0, "errors": errors}


# ─── Build Graph ─────────────────────────────────────

def build_document_extraction_graph():
    graph = StateGraph(DocumentExtractionState)

    graph.add_node("parse_document", parse_document_node)
    graph.add_node("extract_structure", extract_structure_node)
    graph.add_node("validate_output", validate_output_node)

    graph.set_entry_point("parse_document")
    graph.add_edge("parse_document", "extract_structure")
    graph.add_edge("extract_structure", "validate_output")
    graph.add_edge("validate_output", END)

    return graph.compile()


document_extraction_agent = build_document_extraction_graph()
